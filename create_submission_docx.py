"""Create Word submission document matching the Canvas discussion board format."""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parent
METRICS_PATH = ROOT / "results" / "metrics.json"
OUTPUT_PATH = ROOT / "results" / "Financial_RAG_Discussion_Board_Submission.docx"

STUDENT_NAME = "Audrey Rah"
GITHUB_LINK = "https://github.com/YOUR_GITHUB_USERNAME/financial-rag-challenge"


def fmt_pct(value: float) -> str:
    return f"{value:.2f}%"


def fmt_mrr(value: float) -> str:
    return f"{value:.4f}"


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)


def main() -> None:
    metrics = json.loads(METRICS_PATH.read_text(encoding="utf-8"))
    baseline = metrics["baseline"]
    engineered = metrics["engineered"]
    years = ", ".join(str(y) for y in metrics["years"])

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_paragraph(f"Name: {STUDENT_NAME} | Recent Years Used: {years}")
    doc.add_paragraph(f"Github Link to the work: {GITHUB_LINK}")

    add_heading(doc, "Part 1: The Scorecard")
    doc.add_paragraph()

    table = doc.add_table(rows=6, cols=3)
    table.style = "Table Grid"
    headers = ["Metric", "Baseline (Simple)", "Engineered (Improved)"]
    for col, header in enumerate(headers):
        cell = table.rows[0].cells[col]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    rows = [
        ("Hit Rate (K=5)", fmt_pct(baseline["hit_rate_at_k"]), fmt_pct(engineered["hit_rate_at_k"])),
        ("MRR", fmt_mrr(baseline["mrr"]), fmt_mrr(engineered["mrr"])),
        ("Groundedness", fmt_pct(baseline["groundedness"]), fmt_pct(engineered["groundedness"])),
        ("Factual Accuracy", fmt_pct(baseline["factual_accuracy"]), fmt_pct(engineered["factual_accuracy"])),
        ("Hallucination Rate", fmt_pct(baseline["hallucination_rate"]), fmt_pct(engineered["hallucination_rate"])),
    ]
    for row_idx, (metric, base_val, eng_val) in enumerate(rows, start=1):
        table.rows[row_idx].cells[0].text = metric
        table.rows[row_idx].cells[1].text = base_val
        table.rows[row_idx].cells[2].text = eng_val

    for row in table.rows:
        row.height = Inches(0.28)

    doc.add_paragraph()
    add_heading(doc, "Part 2: Engineering Reflection")
    doc.add_paragraph()

    doc.add_paragraph(
        "The Bottleneck: Looking at my Baseline, the main failure was in Finding the data (Retriever), "
        "not Understanding the data (Generator). Hit Rate@5 was only "
        f"{fmt_pct(baseline['hit_rate_at_k'])} and MRR was {fmt_mrr(baseline['mrr'])}, while Groundedness "
        f"stayed at {fmt_pct(baseline['groundedness'])}. That gap shows the librarian often returned the wrong "
        "bulletin in the top 5, even though any number it picked usually came from the retrieved text. "
        f"Factual Accuracy at {fmt_pct(baseline['factual_accuracy'])} confirms the system rarely produced the "
        "correct CSV answer — but the low Hit Rate and MRR pointed to retrieval as the primary bottleneck."
    )

    doc.add_paragraph(
        "The Metadata Fix: Adding Year/Month metadata changed scores mainly on the retrieval side. After tagging "
        "each chunk with year and month from filenames like treasury_bulletin_2022_12.txt, and applying year-level "
        "filtering plus soft rerank bonuses in the engineered pipeline, Hit Rate@5 rose from "
        f"{fmt_pct(baseline['hit_rate_at_k'])} to {fmt_pct(engineered['hit_rate_at_k'])} and MRR from "
        f"{fmt_mrr(baseline['mrr'])} to {fmt_mrr(engineered['mrr'])}. Factual Accuracy improved from "
        f"{fmt_pct(baseline['factual_accuracy'])} to {fmt_pct(engineered['factual_accuracy'])}. Groundedness stayed "
        f"at {fmt_pct(engineered['groundedness'])} because the engineered generator only outputs verbatim context "
        "numbers or table math backed by source operands. Metadata helped retrieval metrics more than generation metrics."
    )

    doc.add_paragraph(
        "Scaling Insight: If I scaled from this 4-year subset to the full 80-year archive (1939–2025), the first "
        "component that would likely break is the in-memory vector search layer (embedding matrix + cosine similarity "
        "over all chunks). Chunk count grows with years, so brute-force nearest-neighbor search and full-corpus "
        "re-embedding would become too slow and memory-heavy long before the answer generator became the bottleneck."
    )

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Saved {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
